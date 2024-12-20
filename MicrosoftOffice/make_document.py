#!/usr/bin/env python3
"""Generate a Word document from input data."""

import argparse
import csv
import random
from pathlib import Path
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from PIL import Image, ImageDraw


def generate_abstract_image(
    width: int = 400, height: int = 400, num_circles: int = 50, seed="abstract"
):
    """Generate a random image."""
    # Create a blank white image
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    random.seed(seed)

    # Generate random colorful circles
    for _ in range(num_circles):
        # Random circle position and size
        x1 = random.randint(0, width)
        y1 = random.randint(0, height)
        radius = random.randint(20, 80)
        x2 = x1 + radius
        y2 = y1 + radius

        # Random color (RGBA with some transparency)
        fill_color = (
            random.randint(0, 255),  # Red
            random.randint(0, 255),  # Green
            random.randint(0, 255),  # Blue
            200,  # Alpha for transparency
        )

        # Draw the circle
        draw.ellipse([x1, y1, x2, y2], fill=fill_color[:3], outline=None)

    buffer = BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def create_word_document(csv_file: Path, outname: Optional[str] = None):
    """Create a word document."""
    doc = Document()

    # Add a title
    doc.add_heading("Demonstration of python-docx Features", level=1)

    # Add a paragraph with different formatting
    para = doc.add_paragraph("This is a paragraph with some ")
    bold_text = para.add_run("bold")
    bold_text.bold = True
    para.add_run(" and ")
    italic_text = para.add_run("italic")
    italic_text.italic = True
    para.add_run(" text.")

    # Add another paragraph with custom style
    style = doc.styles.add_style("CustomStyle", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Arial"
    style.font.size = Pt(14)
    style.font.bold = True
    doc.add_paragraph("This paragraph uses a custom style.", style="CustomStyle")

    # Add a heading
    doc.add_heading("Tables", level=2)

    # Read the CSV file
    with open(csv_file, "r", newline="", encoding="utf-8") as file:
        reader = csv.DictReader(file)
        rows = list(reader)

    # Add a table
    table = doc.add_table(rows=len(rows) + 1, cols=len(rows[0]))
    table.style = "Table Grid"

    for col_idx, column_name in enumerate(rows[0].keys()):
        table.rows[0].cells[col_idx].text = column_name

    for row, table_row in zip(rows, table.rows[1:]):
        for cell, value in zip(table_row.cells, row.values()):
            cell.text = value

    # Add a heading
    doc.add_heading("Images", level=2)

    # Add an image
    doc.add_paragraph("Below is an example image:")
    doc.add_picture(generate_abstract_image(), width=Inches(2), height=Inches(2))

    # Add a section with a header and footer
    section = doc.sections[-1]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Header text"
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Footer text"

    # Add a page break
    doc.add_page_break()

    # Add a bulleted list
    doc.add_heading("Lists", level=2)
    doc.add_paragraph("Item 1", style="List Bullet")
    doc.add_paragraph("Item 2", style="List Bullet")
    doc.add_paragraph("Item 3", style="List Bullet")

    # Add a numbered list
    doc.add_paragraph("Step 1", style="List Number")
    doc.add_paragraph("Step 2", style="List Number")
    doc.add_paragraph("Step 3", style="List Number")

    # Save the document
    output_path = csv_file.with_suffix(".docx")
    if outname:
        output_path = output_path.with_stem(outname)

    doc.save(str(output_path))

    print(f"Saved document as {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("csvfile", type=Path)
    parser.add_argument("--outname", type=str)

    args = parser.parse_args()
    create_word_document(args.csvfile, args.outname)
